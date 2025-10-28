
"""
문서 파일 파서
Word, PDF, Markdown 파일을 텍스트로 변환
"""

from typing import Optional
import docx  # python-docx
try:
    import pdfplumber
except Exception:
    pdfplumber = None
import PyPDF2
import markdown
from pathlib import Path
import re


class DocumentParser:
    """문서 파일 파서"""
    
    SUPPORTED_FORMATS = ['.docx', '.pdf', '.md']
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
    
    @staticmethod
    def parse_file(file_path: str) -> str:
        """
        파일 형식에 맞는 파서 선택 및 텍스트 추출
        
        Args:
            file_path: 파일 경로
            
        Returns:
            추출된 텍스트
            
        Raises:
            ValueError: 지원하지 않는 파일 형식
            IOError: 파일 읽기 오류
        """
        path = Path(file_path)
        
        # 파일 존재 확인
        if not path.exists():
            raise FileNotFoundError(f"파일을 찾을 수 없습니다: {file_path}")
        
        # 파일 크기 확인
        if path.stat().st_size > DocumentParser.MAX_FILE_SIZE:
            raise ValueError(f"파일 크기가 너무 큽니다 (최대 10MB)")
        
        # 확장자별 파싱
        ext = path.suffix.lower()
        
        if ext == '.docx':
            return DocumentParser._parse_docx(file_path)
        elif ext == '.pdf':
            return DocumentParser._parse_pdf(file_path)
        elif ext == '.md':
            return DocumentParser._parse_markdown(file_path)
        else:
            raise ValueError(f"지원하지 않는 파일 형식: {ext}")
    
    @staticmethod
    def _parse_docx(file_path: str) -> str:
        """Word 문서 파싱"""
        try:
            doc = docx.Document(file_path)
            
            # 모든 단락 추출
            paragraphs = [para.text for para in doc.paragraphs]
            
            # 표 내용 추출
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text for cell in row.cells)
                    tables_text.append(row_text)
            
            # 합치기
            full_text = '\n'.join(paragraphs)
            if tables_text:
                full_text += '\n\n[표 데이터]\n' + '\n'.join(tables_text)
            
            return full_text.strip()
            
        except Exception as e:
            raise IOError(f"Word 문서 파싱 실패: {str(e)}")
    
    @staticmethod
    def _parse_pdf(file_path: str) -> str:
        """PDF 문서 파싱"""
        try:
            # 우선 pdfplumber 시도 (한글/레이아웃 공백 유지에 유리)
            if pdfplumber is not None:
                text_parts = []
                with pdfplumber.open(file_path) as pdf:
                    for i, page in enumerate(pdf.pages, 1):
                        # tolerance 조정으로 공백 보존 개선
                        page_text = page.extract_text(x_tolerance=1.5, y_tolerance=1.0) or ''
                        page_text = DocumentParser._normalize_text(page_text)
                        if page_text.strip():
                            text_parts.append(f"[Page {i}]\n{page_text}")
                if text_parts:
                    return '\n\n'.join(text_parts).strip()

            # 폴백: PyPDF2
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_parts = []
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    page_text = page.extract_text() or ''
                    page_text = DocumentParser._normalize_text(page_text)
                    if page_text.strip():
                        text_parts.append(f"[Page {page_num}]\n{page_text}")
                return '\n\n'.join(text_parts).strip()
        except Exception as e:
            raise IOError(f"PDF 문서 파싱 실패: {str(e)}")

    @staticmethod
    def _normalize_text(text: str) -> str:
        """PDF 추출 텍스트 정규화
        - 유니코드 정규화(NFC)
        - 제어문자 제거
        - 불릿 문자 통일(•, ◦ 등 → •)
        - 라인/공백 정리
        """
        try:
            import unicodedata, re
            s = text or ''
            s = unicodedata.normalize('NFC', s)
            # 제어문자 제거
            s = ''.join(ch for ch in s if not unicodedata.category(ch).startswith('C'))
            # 불릿 통일
            s = s.replace('◦', '•').replace('\u2022', '•')
            # 줄바꿈 정리: 너무 긴 단어 붙음 방지 위해 줄 단위 공백 하나 삽입
            lines = [ln.strip() for ln in s.splitlines()]
            s = '\n'.join(lines)
            # 여러 공백을 하나로
            s = re.sub(r"\s+", " ", s)
            # 페이지 헤더 간 구분 위해 문장부호 후 공백 보장
            s = re.sub(r"([\.!?])(?=\S)", r"\1 ", s)
            return s
        except Exception:
            return text or ''
    
    @staticmethod
    def _parse_markdown(file_path: str) -> str:
        """Markdown 문서 파싱"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                md_content = file.read()
            
            # Markdown을 일반 텍스트로 변환 (HTML 태그 제거)
            html = markdown.markdown(md_content)
            
            # HTML 태그 제거 (간단한 방법)
            text = re.sub('<[^<]+?>', '', html)
            text = re.sub(r'\n\s*\n', '\n\n', text)  # 중복 개행 제거
            
            return text.strip()
            
        except Exception as e:
            raise IOError(f"Markdown 문서 파싱 실패: {str(e)}")
